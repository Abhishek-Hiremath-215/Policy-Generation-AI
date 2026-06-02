"""
Document processing service for extracting text from various file formats.
"""

import json
import pandas as pd
from pathlib import Path
from typing import Optional

# PDF processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX processing - using docx2txt as fallback
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    try:
        import docx2txt
        DOCX_AVAILABLE = "docx2txt"
    except ImportError:
        DOCX_AVAILABLE = False


class DocumentProcessor:
    """Extracts text from various document formats."""
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        if not PyPDF2:
            raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
        
        try:
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text.append(page_text)
            return "\n\n".join(text)
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        try:
            if DOCX_AVAILABLE == "docx2txt":
                # Use docx2txt as fallback
                import docx2txt
                text = docx2txt.process(file_path)
                return text
            else:
                # Use python-docx
                doc = DocxDocument(file_path)
                text = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = " | ".join([cell.text for cell in row.cells])
                        text.append(row_text)
                
                return "\n\n".join(text)
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_csv(file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            df = pd.read_csv(file_path)
            
            # Create structured text representation
            text_parts = []
            
            # Add column headers
            text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
            text_parts.append("\n")
            
            # Add rows
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                text_parts.append(f"Row {idx + 1}: {row_text}")
            
            # Add summary statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("\n\nSummary Statistics:")
                text_parts.append(df[numeric_cols].describe().to_string())
            
            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting CSV: {str(e)}")
    
    @staticmethod
    def extract_text_from_excel(file_path: str) -> str:
        """Extract text from Excel file."""
        try:
            # Read all sheets
            excel_file = pd.ExcelFile(file_path)
            text_parts = []
            
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                text_parts.append(f"\n\n=== Sheet: {sheet_name} ===")
                text_parts.append("Columns: " + ", ".join(df.columns.tolist()))
                text_parts.append("\n")
                
                # Add rows
                for idx, row in df.iterrows():
                    row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                    text_parts.append(f"Row {idx + 1}: {row_text}")
                
                # Add summary for numeric columns
                numeric_cols = df.select_dtypes(include=['number']).columns
                if len(numeric_cols) > 0:
                    text_parts.append("\nSummary Statistics:")
                    text_parts.append(df[numeric_cols].describe().to_string())
            
            return "\n".join(text_parts)
        except Exception as e:
            raise Exception(f"Error extracting Excel: {str(e)}")
    
    @staticmethod
    def extract_text_from_json(file_path: str) -> str:
        """Extract text from JSON file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                data = json.load(file)
            
            # Convert JSON to readable text format
            def json_to_text(obj, prefix=""):
                text_parts = []
                
                if isinstance(obj, dict):
                    for key, value in obj.items():
                        if isinstance(value, (dict, list)):
                            text_parts.append(f"{prefix}{key}:")
                            text_parts.extend(json_to_text(value, prefix + "  "))
                        else:
                            text_parts.append(f"{prefix}{key}: {value}")
                
                elif isinstance(obj, list):
                    for idx, item in enumerate(obj):
                        if isinstance(item, (dict, list)):
                            text_parts.append(f"{prefix}[{idx}]:")
                            text_parts.extend(json_to_text(item, prefix + "  "))
                        else:
                            text_parts.append(f"{prefix}[{idx}]: {item}")
                
                else:
                    text_parts.append(f"{prefix}{obj}")
                
                return text_parts
            
            return "\n".join(json_to_text(data))
        except Exception as e:
            raise Exception(f"Error extracting JSON: {str(e)}")
    
    @staticmethod
    def process_document(file_path: str, file_type: str) -> str:
        """
        Process document based on file type.
        
        Args:
            file_path: Path to the document
            file_type: Type of document (pdf, docx, csv, etc.)
            
        Returns:
            Extracted text content
        """
        extractors = {
            'pdf': DocumentProcessor.extract_text_from_pdf,
            'docx': DocumentProcessor.extract_text_from_docx,
            'csv': DocumentProcessor.extract_text_from_csv,
            'xlsx': DocumentProcessor.extract_text_from_excel,
            'xls': DocumentProcessor.extract_text_from_excel,
            'json': DocumentProcessor.extract_text_from_json,
        }
        
        extractor = extractors.get(file_type)
        if not extractor:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        return extractor(file_path)
